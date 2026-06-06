import os
import sys
import django
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Setup Django Environment
backend_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'animacare_backend')
sys.path.append(backend_dir)
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from django.apps import apps

def generate_docx():
    doc = Document()
    
    # Optional: Change default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    models = apps.get_models()
    
    for model in models:
        # Only process models from our apps
        if not model.__module__.startswith('apps.'):
            continue
            
        table_name = model.__name__
        primary_key = []
        foreign_keys = []
        fields_data = []
        
        # Ensure we have an ID field if not explicitly defined
        pk_field = model._meta.pk
        
        for idx, field in enumerate(model._meta.fields, 1):
            col_name = field.name
            
            # Primary Key
            is_pk = (field == pk_field)
            if is_pk:
                primary_key.append(col_name)
                
            # Foreign Key
            is_fk = field.is_relation and (field.many_to_one or field.one_to_one)
            if is_fk:
                foreign_keys.append(col_name)
                
            # Datatype & Size
            internal_type = field.get_internal_type()
            datatype = internal_type
            size = ""
            
            if internal_type == 'CharField':
                datatype = 'Varchar'
                size = str(field.max_length) if field.max_length else ""
            elif internal_type in ('IntegerField', 'AutoField', 'BigAutoField', 'PositiveIntegerField'):
                datatype = 'Integer'
                size = "11"
            elif internal_type == 'DateTimeField':
                datatype = 'Datetime'
            elif internal_type == 'DateField':
                datatype = 'Date'
            elif internal_type == 'TextField':
                datatype = 'Text'
            elif internal_type == 'BooleanField':
                datatype = 'Boolean'
            elif internal_type == 'DecimalField':
                datatype = 'Decimal'
                size = f"{field.max_digits},{field.decimal_places}"
            elif internal_type in ('ForeignKey', 'OneToOneField'):
                datatype = 'Integer'
                size = "11"
            elif internal_type == 'UUIDField':
                datatype = 'UUID'
                size = "36"
            elif internal_type == 'URLField':
                datatype = 'Varchar'
                size = str(field.max_length) if field.max_length else "200"
            elif internal_type == 'JSONField':
                datatype = 'JSON'
            elif internal_type == 'FloatField':
                datatype = 'Float'
            elif internal_type == 'GenericIPAddressField':
                datatype = 'Varchar'
                size = "39"
                
            # Constraints
            constraints = []
            if is_pk:
                constraints.append("Primary Key")
                if internal_type in ('AutoField', 'BigAutoField'):
                    constraints.append("Auto Increment")
            if is_fk:
                constraints.append("Foreign Key")
            if field.unique:
                constraints.append("Unique")
            if not field.null and not is_pk:
                constraints.append("Not Null")
                
            # Description
            desc = f"Stores {col_name.replace('_', ' ')}"
            if is_fk:
                desc = f"Stores reference ID to the {field.related_model.__name__}"
            if is_pk:
                desc = f"Stores unique {table_name.lower()} id"
                
            fields_data.append([
                str(idx),
                col_name,
                datatype,
                size,
                ", ".join(constraints),
                desc
            ])
            
        # Write to DOCX
        p1 = doc.add_paragraph()
        r1 = p1.add_run(f"Table Name: {table_name}")
        r1.bold = True
        r1.font.size = Pt(14)
        
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"Primary Key: {', '.join(primary_key) if primary_key else 'id'}")
        r2.bold = True
        r2.font.size = Pt(12)
        
        p3 = doc.add_paragraph()
        fk_str = ', '.join(foreign_keys) if foreign_keys else 'None'
        r3 = p3.add_run(f"Foreign Key: {fk_str}")
        r3.bold = True
        r3.font.size = Pt(12)
        
        # Add table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Make table 100% width
        tbl_pr = table._element.xpath('w:tblPr')
        if tbl_pr:
            tbl_w = OxmlElement('w:tblW')
            tbl_w.set(qn('w:w'), '5000')
            tbl_w.set(qn('w:type'), 'pct')
            tbl_pr[0].append(tbl_w)
        
        # Set column headers
        headers = ["S.NO.", "COLUMN", "DATATYPE", "SIZE", "CONSTRAINTS", "DESCRIPTION"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    
        # Populate table data
        for row_data in fields_data:
            row_cells = table.add_row().cells
            for i, text in enumerate(row_data):
                row_cells[i].text = text
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        
        doc.add_paragraph("\n") # Blank line between models

    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Database_Tables.docx")
    doc.save(output_path)
    print(f"Generated successfully at: {output_path}")

if __name__ == '__main__':
    generate_docx()
